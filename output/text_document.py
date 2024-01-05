import typing

import docx
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

from common import models
from common import settings

LAYOUT = {
    'top': Mm(115), 'bottom': Mm(20), 'left': Mm(40), 'right': Mm(40),  # границы листа
    'skip': Mm(40),  # отступ перед строками с подписями
    'tab_stop': Mm(100),  # позиция табуляции для фамилий подписантов
}

# Стили параграфов диплома
DEFAULT_FONT, DEFAULT_ALIGN = ('Times New Roman', WD_ALIGN_PARAGRAPH.CENTER)
HEAD_STYLE = {'font': DEFAULT_FONT, 'size': Pt(24), 'bold': True, 'italic': True, 'align': DEFAULT_ALIGN}
NAME_STYLE = {'font': DEFAULT_FONT, 'size': Pt(24), 'bold': True, 'italic': True, 'align': DEFAULT_ALIGN}
SCHOOL_STYLE = {'font': DEFAULT_FONT, 'size': Pt(16), 'bold': True, 'italic': True, 'align': DEFAULT_ALIGN}
STATUS_STYLE = {'font': DEFAULT_FONT, 'size': Pt(20), 'bold': True, 'italic': True, 'align': DEFAULT_ALIGN}
SEASON_STYLE = {'font': DEFAULT_FONT, 'size': Pt(16), 'bold': True, 'italic': True, 'align': DEFAULT_ALIGN}
SIGN_STYLE = {'font': DEFAULT_FONT, 'size': Pt(14), 'bold': False, 'italic': False, 'align': WD_ALIGN_PARAGRAPH.LEFT}


class TextDocument:
    def __init__(self, filename: typing.Optional[str] = None):
        self.doc = docx.Document()
        section = self.doc.sections[0]
        section.page_height = Mm(297)  # A4 портрет
        section.page_width = Mm(210)
        section.left_margin = LAYOUT['left']
        section.right_margin = LAYOUT['right']
        section.top_margin = LAYOUT['top']
        section.bottom_margin = LAYOUT['bottom']

        self.filename = filename or settings.RESULT_FILENAME

    def generate(self, diplomas: typing.List[models.Diploma]):
        for numPage, diploma in enumerate(diplomas):
            self.doc.add_page_break() if numPage > 0 else None
            self._make_page(diploma)
        self.doc.save(self.filename)

    def _make_paragraph(self, text, style) -> Paragraph:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = style['align']
        font = paragraph.add_run(text).font
        font.name, font.size, font.bold, font.italic = (
            style['font'], style['size'], style['bold'], style['italic']
        )
        return paragraph

    def _make_page(self, diploma) -> None:
        self._make_paragraph(f'Дипломом {diploma.degree.degree_text}\nнаграждается', HEAD_STYLE)
        self._make_paragraph(diploma.contestant_full_name, NAME_STYLE)
        self._make_paragraph(f'{diploma.school_name}', SCHOOL_STYLE)
        self._make_paragraph(f'{diploma.degree.status_text} финала', STATUS_STYLE)
        self._make_paragraph(f'сезона {settings.SEASON} чемпионата ИТ-сферы Ульяновской области', SEASON_STYLE)
        fmt = self._make_paragraph(settings.SIGN_ULIVT, SIGN_STYLE).paragraph_format
        fmt.space_before = LAYOUT['skip']
        fmt.tab_stops.add_tab_stop(LAYOUT['tab_stop'])
        self._make_paragraph(settings.SIGN_DEAN, SIGN_STYLE).paragraph_format.tab_stops.add_tab_stop(LAYOUT['tab_stop'])
